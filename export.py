"""
Export Manager for TestBuddy

Handles exporting OCR results to multiple formats:
- PDF (with optional image + text overlay)
- DOCX (Microsoft Word format with formatting)
- TXT (plain text)
- CSV (for structured data/tables)
- HTML (web viewable)
- Markdown (documentation format)
- JSON (structured data with metadata)
"""

from __future__ import annotations

import json
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any
from io import BytesIO

# PDF generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
except ImportError:
    SimpleDocTemplate = None

# DOCX generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

# Markdown
try:
    import markdown
except ImportError:
    markdown = None

# PIL for image handling
try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None


class ExportManager:
    """Manages exporting OCR results to various formats."""
    
    def __init__(self, export_directory: Path = None) -> None:
        """Initialize export manager.
        
        Args:
            export_directory: Directory to save exports. Defaults to ./exports/
        """
        self.export_dir = export_directory or Path.cwd() / "exports"
        self.export_dir.mkdir(exist_ok=True)
    
    def _sanitize_filename(self, filename: str) -> str:
        """Remove invalid filename characters.
        
        Args:
            filename: Original filename
            
        Returns:
            Sanitized filename
        """
        invalid_chars = r'<>:"/\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, "_")
        return filename[:255]  # Filename length limit
    
    def _get_export_path(self, name: str, extension: str) -> Path:
        """Get full export file path.
        
        Args:
            name: Session/document name
            extension: File extension (pdf, docx, txt, etc)
            
        Returns:
            Full Path to export file
        """
        safe_name = self._sanitize_filename(name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{safe_name}_{timestamp}.{extension}"
        return self.export_dir / filename
    
    # ============= TEXT EXPORT =============
    
    def to_txt(self, text: str, name: str, metadata: Optional[Dict] = None) -> Path:
        """Export to plain text file.
        
        Args:
            text: OCR text content
            name: Session name for filename
            metadata: Optional metadata dict (language, category, tags)
            
        Returns:
            Path to exported file
        """
        export_path = self._get_export_path(name, "txt")
        
        # Prepare content
        content = text
        
        # Add metadata header if provided
        if metadata:
            header = "=== TESTBUDDY OCR EXPORT ===\n"
            if metadata.get("session_name"):
                header += f"Session: {metadata['session_name']}\n"
            if metadata.get("language"):
                header += f"Language: {metadata['language']}\n"
            if metadata.get("category"):
                header += f"Category: {metadata['category']}\n"
            if metadata.get("tags"):
                header += f"Tags: {', '.join(metadata['tags'])}\n"
            if metadata.get("created_at"):
                header += f"Created: {metadata['created_at']}\n"
            header += "=" * 28 + "\n\n"
            content = header + text
        
        # Write file
        with open(export_path, "w", encoding="utf-8") as f:
            f.write(content)
        
        return export_path
    
    # ============= PDF EXPORT =============
    
    def to_pdf(self, text: str, name: str, image: Optional[Any] = None, 
               metadata: Optional[Dict] = None) -> Path:
        """Export to PDF with optional image.
        
        Args:
            text: OCR text content
            name: Session name for filename
            image: Optional PIL Image or file path
            metadata: Optional metadata dict
            
        Returns:
            Path to exported file
        """
        if not SimpleDocTemplate:
            raise ImportError("reportlab not installed. Install with: pip install reportlab")
        
        export_path = self._get_export_path(name, "pdf")
        
        # Create PDF
        doc = SimpleDocTemplate(
            str(export_path),
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        # Prepare styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=12
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            leading=14
        )
        
        # Build story
        story = []
        
        # Add title
        story.append(Paragraph(name or "OCR Result", title_style))
        
        # Add metadata
        if metadata:
            meta_text = []
            if metadata.get("language"):
                meta_text.append(f"<i>Language: {metadata['language']}</i>")
            if metadata.get("created_at"):
                meta_text.append(f"<i>Created: {metadata['created_at']}</i>")
            if meta_text:
                story.append(Paragraph(" | ".join(meta_text), body_style))
            story.append(Spacer(1, 12))
        
        # Add image if provided
        if image:
            try:
                if isinstance(image, str):
                    img = RLImage(image, width=5.5*inch, height=4*inch)
                elif isinstance(image, PILImage.Image):
                    img_bytes = BytesIO()
                    image.save(img_bytes, format="PNG")
                    img_bytes.seek(0)
                    img = RLImage(img_bytes, width=5.5*inch, height=4*inch)
                else:
                    img = None
                
                if img:
                    story.append(img)
                    story.append(Spacer(1, 12))
            except Exception as e:
                print(f"Warning: Could not embed image in PDF: {e}")
        
        # Add text content
        story.append(Paragraph("<b>OCR Text:</b>", styles['Heading2']))
        story.append(Spacer(1, 6))
        
        # Format text (preserve paragraphs)
        for para in text.split("\n"):
            if para.strip():
                story.append(Paragraph(para, body_style))
            else:
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        return export_path
    
    # ============= DOCX EXPORT =============
    
    def to_docx(self, text: str, name: str, image: Optional[Any] = None,
                metadata: Optional[Dict] = None) -> Path:
        """Export to DOCX (Microsoft Word format).
        
        Args:
            text: OCR text content
            name: Session name for filename
            image: Optional PIL Image or file path
            metadata: Optional metadata dict
            
        Returns:
            Path to exported file
        """
        if not Document:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        export_path = self._get_export_path(name, "docx")
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading(name or "OCR Result", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        if metadata:
            meta_para = doc.add_paragraph()
            if metadata.get("language"):
                meta_para.add_run(f"Language: {metadata['language']} | ")
            if metadata.get("created_at"):
                meta_para.add_run(f"Created: {metadata['created_at']}")
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_para.runs[0].font.size = Pt(9)
            meta_para.runs[0].font.italic = True
            doc.add_paragraph()  # Blank line
        
        # Add image if provided
        if image:
            try:
                if isinstance(image, str):
                    doc.add_picture(image, width=Inches(5.5))
                elif isinstance(image, PILImage.Image):
                    img_bytes = BytesIO()
                    image.save(img_bytes, format="PNG")
                    img_bytes.seek(0)
                    doc.add_picture(img_bytes, width=Inches(5.5))
                doc.add_paragraph()  # Blank line
            except Exception as e:
                print(f"Warning: Could not embed image in DOCX: {e}")
        
        # Add text content
        doc.add_heading("OCR Text", level=2)
        
        # Format text (preserve paragraphs)
        for para in text.split("\n"):
            if para.strip():
                p = doc.add_paragraph(para)
                p.paragraph_format.line_spacing = 1.15
            else:
                doc.add_paragraph()
        
        # Save document
        doc.save(str(export_path))
        return export_path
    
    # ============= CSV EXPORT =============
    
    def to_csv(self, text: str, name: str, metadata: Optional[Dict] = None) -> Path:
        """Export to CSV format (useful for structured data/tables).
        
        Args:
            text: OCR text content
            name: Session name for filename
            metadata: Optional metadata dict
            
        Returns:
            Path to exported file
        """
        import csv
        
        export_path = self._get_export_path(name, "csv")
        
        # For simple text, create lines CSV
        lines = text.split("\n")
        
        with open(export_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            
            # Write metadata as header rows if provided
            if metadata:
                writer.writerow(["Session", name or ""])
                if metadata.get("language"):
                    writer.writerow(["Language", metadata["language"]])
                if metadata.get("created_at"):
                    writer.writerow(["Created", metadata["created_at"]])
                if metadata.get("tags"):
                    writer.writerow(["Tags", ", ".join(metadata["tags"])])
                writer.writerow([])  # Blank row
            
            # Write content lines
            writer.writerow(["Line Number", "Text"])
            for i, line in enumerate(lines, 1):
                writer.writerow([i, line])
        
        return export_path
    
    # ============= HTML EXPORT =============
    
    def to_html(self, text: str, name: str, image: Optional[Any] = None,
                metadata: Optional[Dict] = None) -> Path:
        """Export to HTML (web viewable).
        
        Args:
            text: OCR text content
            name: Session name for filename
            image: Optional PIL Image or file path
            metadata: Optional metadata dict
            
        Returns:
            Path to exported file
        """
        export_path = self._get_export_path(name, "html")
        
        # Create HTML content
        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            "<meta charset='UTF-8'>",
            f"<title>{name or 'OCR Result'}</title>",
            "<style>",
            "body { font-family: Arial, sans-serif; margin: 2em; max-width: 800px; }",
            "h1 { color: #1f77b4; border-bottom: 2px solid #1f77b4; }",
            ".metadata { color: #666; font-size: 0.9em; margin: 1em 0; }",
            ".content { line-height: 1.6; white-space: pre-wrap; }",
            "img { max-width: 100%; border: 1px solid #ddd; margin: 1em 0; }",
            "</style>",
            "</head>",
            "<body>",
        ]
        
        # Add title
        html_parts.append(f"<h1>{name or 'OCR Result'}</h1>")
        
        # Add metadata
        if metadata:
            html_parts.append("<div class='metadata'>")
            if metadata.get("language"):
                html_parts.append(f"<p><strong>Language:</strong> {metadata['language']}</p>")
            if metadata.get("created_at"):
                html_parts.append(f"<p><strong>Created:</strong> {metadata['created_at']}</p>")
            if metadata.get("tags"):
                html_parts.append(f"<p><strong>Tags:</strong> {', '.join(metadata['tags'])}</p>")
            html_parts.append("</div>")
        
        # Add image if provided (embedded as base64 or linked)
        if image:
            try:
                if isinstance(image, str):
                    html_parts.append(f"<img src='{image}' alt='OCR Image'>")
                elif isinstance(image, PILImage.Image):
                    import base64
                    img_bytes = BytesIO()
                    image.save(img_bytes, format="PNG")
                    img_bytes.seek(0)
                    b64_str = base64.b64encode(img_bytes.getvalue()).decode()
                    html_parts.append(f"<img src='data:image/png;base64,{b64_str}' alt='OCR Image'>")
            except Exception as e:
                print(f"Warning: Could not embed image in HTML: {e}")
        
        # Add text content
        html_parts.append("<div class='content'>")
        html_parts.append(text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
        html_parts.append("</div>")
        
        # Close HTML
        html_parts.append("</body>")
        html_parts.append("</html>")
        
        # Write file
        with open(export_path, "w", encoding="utf-8") as f:
            f.write("\n".join(html_parts))
        
        return export_path
    
    # ============= MARKDOWN EXPORT =============
    
    def to_markdown(self, text: str, name: str, metadata: Optional[Dict] = None) -> Path:
        """Export to Markdown format.
        
        Args:
            text: OCR text content
            name: Session name for filename
            metadata: Optional metadata dict
            
        Returns:
            Path to exported file
        """
        export_path = self._get_export_path(name, "md")
        
        md_parts = [f"# {name or 'OCR Result'}\n"]
        
        # Add metadata
        if metadata:
            md_parts.append("## Document Information\n")
            if metadata.get("language"):
                md_parts.append(f"- **Language:** {metadata['language']}\n")
            if metadata.get("created_at"):
                md_parts.append(f"- **Created:** {metadata['created_at']}\n")
            if metadata.get("category"):
                md_parts.append(f"- **Category:** {metadata['category']}\n")
            if metadata.get("tags"):
                md_parts.append(f"- **Tags:** {', '.join(metadata['tags'])}\n")
            md_parts.append("\n")
        
        # Add content
        md_parts.append("## OCR Content\n\n")
        md_parts.append(text)
        
        # Write file
        with open(export_path, "w", encoding="utf-8") as f:
            f.write("".join(md_parts))
        
        return export_path
    
    # ============= JSON EXPORT =============
    
    def to_json(self, text: str, name: str, metadata: Optional[Dict] = None,
                image_path: Optional[str] = None) -> Path:
        """Export to JSON with full metadata and structure.
        
        Args:
            text: OCR text content
            name: Session name for filename
            metadata: Optional metadata dict
            image_path: Optional path to source image
            
        Returns:
            Path to exported file
        """
        export_path = self._get_export_path(name, "json")
        
        data = {
            "session_name": name,
            "export_timestamp": datetime.now().isoformat(),
            "text": text,
            "text_length": len(text),
            "text_lines": len(text.split("\n")),
            "text_words": len(text.split()),
            "metadata": metadata or {},
            "image_path": image_path
        }
        
        with open(export_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        return export_path
    
    # ============= MULTI-FORMAT EXPORT =============
    
    def export_all_formats(self, text: str, name: str, image: Optional[Any] = None,
                          metadata: Optional[Dict] = None,
                          formats: Optional[list] = None) -> Dict[str, Path]:
        """Export to multiple formats at once.
        
        Args:
            text: OCR text content
            name: Session name
            image: Optional image
            metadata: Optional metadata
            formats: List of formats (pdf, docx, txt, csv, html, md, json)
                    Defaults to all formats
            
        Returns:
            Dict mapping format to export Path
        """
        if not formats:
            formats = ["txt", "pdf", "docx", "html", "json"]
        
        results = {}
        
        for fmt in formats:
            try:
                if fmt.lower() == "pdf":
                    results["pdf"] = self.to_pdf(text, name, image, metadata)
                elif fmt.lower() == "docx":
                    results["docx"] = self.to_docx(text, name, image, metadata)
                elif fmt.lower() == "txt":
                    results["txt"] = self.to_txt(text, name, metadata)
                elif fmt.lower() == "csv":
                    results["csv"] = self.to_csv(text, name, metadata)
                elif fmt.lower() == "html":
                    results["html"] = self.to_html(text, name, image, metadata)
                elif fmt.lower() == "md":
                    results["md"] = self.to_markdown(text, name, metadata)
                elif fmt.lower() == "json":
                    results["json"] = self.to_json(text, name, metadata)
            except Exception as e:
                print(f"Error exporting to {fmt}: {e}")
        
        return results
